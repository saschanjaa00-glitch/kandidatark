"""
Student Exam Information - Excel to Word Converter
Reads student exam data from Excel and generates individual A4 pages in Word format
"""

import pandas as pd
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os
import sys


def add_horizontal_line(paragraph):
    """Add a horizontal line below a paragraph"""
    p = paragraph._element
    pPr = p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '12')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '000000')
    pBdr.append(bottom)
    pPr.append(pBdr)


def read_excel_file(file_path):
    """Read Excel file and return DataFrame with student information"""
    try:
        df = pd.read_excel(file_path)
        
        # Check if all required columns are present
        required_columns = [
            'Rom', 'Fagkode', 'Fagnavn', 'Fornavn', 
            'Etternavn', 'Eksamensdato', 'PAS kandidatnummer'
        ]
        
        missing_columns = [col for col in required_columns if col not in df.columns]
        if missing_columns:
            print(f"Warning: Missing columns: {', '.join(missing_columns)}")
            print(f"Available columns: {', '.join(df.columns)}")
            return None
            
        return df
    except Exception as e:
        print(f"Error reading Excel file: {e}")
        return None


def format_date(date_value):
    """Format date value to string"""
    if pd.isna(date_value):
        return ""
    if isinstance(date_value, str):
        return date_value
    try:
        return date_value.strftime("%d.%m.%Y")
    except:
        return str(date_value)


def create_student_page(doc, student_data, is_first=False):
    """Create a formatted A4 page for one student"""
    
    if not is_first:
        doc.add_page_break()
    
    # Add some top margin spacing
    doc.add_paragraph()
    
    # Exam Date - Top of page
    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_run = date_para.add_run(f"Eksamensdato: {format_date(student_data['Eksamensdato'])}")
    date_run.font.size = Pt(16)
    date_run.font.bold = True
    
    doc.add_paragraph()
    
    # Room Number
    room_para = doc.add_paragraph()
    room_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    room_run = room_para.add_run(f"Rom: {student_data['Rom']}")
    room_run.font.size = Pt(20)
    room_run.font.bold = True
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Student Name
    name_para = doc.add_paragraph()
    name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_text = f"{student_data['Fornavn']} {student_data['Etternavn']}"
    name_run = name_para.add_run(name_text)
    name_run.font.size = Pt(24)
    name_run.font.bold = True
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Subject Code
    code_para = doc.add_paragraph()
    code_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    code_run = code_para.add_run(f"Fagkode: {student_data['Fagkode']}")
    code_run.font.size = Pt(18)
    code_run.font.bold = True
    
    doc.add_paragraph()
    
    # Subject Name
    subject_para = doc.add_paragraph()
    subject_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subject_run = subject_para.add_run(student_data['Fagnavn'])
    subject_run.font.size = Pt(16)
    
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()
    
    # PAS Kandidatnummer - Most prominent
    pas_para = doc.add_paragraph()
    pas_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pas_label = pas_para.add_run("PAS KANDIDATNUMMER\n")
    pas_label.font.size = Pt(18)
    pas_label.font.bold = True
    
    pas_number = pas_para.add_run(str(student_data['PAS kandidatnummer']))
    pas_number.font.size = Pt(36)
    pas_number.font.bold = True
    pas_number.font.color.rgb = RGBColor(0, 0, 0)


def generate_word_document(df, output_path):
    """Generate Word document with one page per student"""
    doc = Document()
    
    # Set page margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    # Create a page for each student
    for index, row in df.iterrows():
        create_student_page(doc, row, is_first=(index == 0))
    
    # Save the document
    doc.save(output_path)
    print(f"\nDocument created successfully: {output_path}")
    print(f"Total students: {len(df)}")


def main():
    print("=" * 60)
    print("Student Exam Information - Excel to Word Converter")
    print("=" * 60)
    print()
    
    # Get Excel file path from user
    if len(sys.argv) > 1:
        excel_file = sys.argv[1]
    else:
        excel_file = input("Enter the path to the Excel file: ").strip('"').strip("'")
    
    if not os.path.exists(excel_file):
        print(f"Error: File not found: {excel_file}")
        return
    
    # Read Excel file
    print(f"\nReading Excel file: {excel_file}")
    df = read_excel_file(excel_file)
    
    if df is None or df.empty:
        print("Error: Could not read data from Excel file or file is empty.")
        return
    
    print(f"Found {len(df)} students")
    
    # Generate output filename
    output_file = os.path.join(os.path.dirname(excel_file) or '.', 'student_exam_cards.docx')
    
    # Generate Word document
    print(f"\nGenerating Word document...")
    generate_word_document(df, output_file)
    
    print("\nDone!")


if __name__ == "__main__":
    main()
